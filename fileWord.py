import pandas as pd
import docx
import os

variable = docx.Document(r"C:\Users\Asus\Desktop\new.docx")


print(variable.paragraphs[0].text)
p = variable.paragraphs[0];

p.runs[0].bold
p.runs[0].italic
p.runs[0].underline = True

variable.save(r"C:\Users\Asus\Desktop\new.docx")

def getText(fileName):
    doc = docx.Document(fileName)
    fullText = [];

    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText(r"C:\Users\Asus\Desktop\new.docx"))

